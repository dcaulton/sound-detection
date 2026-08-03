import logging
import math
import os
import tempfile
from datetime import UTC, datetime, timedelta
from io import BytesIO
from pathlib import Path
from typing import Any, TypedDict
from uuid import UUID

import requests
from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import Runnable
from langchain_ollama import ChatOllama
from neo4j import Driver
from PIL import Image as PILImage
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlmodel import col

from sound_detection.db.models import BiomeSummary, Detection, Microphone, Recording
from sound_detection.knowledge.rag.retriever import Retriever
from sound_detection.knowledge.species_knowledge_service import SpeciesKnowledgeService

log = logging.getLogger(__name__)
log.setLevel(os.getenv("LOG_LEVEL", "INFO"))

llm = ChatOllama(model="qwen2.5:32b", temperature=0.3)

NON_SPECIES_LABELS = {
    "power tools",
    "fireworks",
    "engine",
    "human vocal",
    "gun",
    "siren",
    "dog",
    "cat",
    "vehicle",
    "airplane",
    "helicopter",
    "chainsaw",
    "lawn mower",
    "human non-vocal",
    # add more as we discover them
}

SPECIES_GROUPS = {
    # Owls
    "bubo": "owl",
    "strix": "owl",
    "megascops": "owl",
    "asio": "owl",
    "aegolius": "owl",
    "otus": "owl",
    "tyto": "owl",
    "surnia": "owl",
    # Raptors (hawks, eagles, falcons, kites, etc.)
    "buteo": "raptor",
    "accipiter": "raptor",
    "haliaeetus": "raptor",
    "circus": "raptor",
    "pandion": "raptor",
    "falco": "raptor",
    "aquila": "raptor",
    "elanus": "raptor",
    "ictinia": "raptor",
    "cathartes": "raptor",  # Turkey Vulture
    "coragyps": "raptor",  # Black Vulture
    # Hummingbirds
    "archilochus": "hummingbird",
    "selasphorus": "hummingbird",
    "calypte": "hummingbird",
    "amazilia": "hummingbird",
    "cynanthus": "hummingbird",
}

HIGH_INTEREST_SPECIES = {
    "tympanuchus cupido",
    "tympanuchus pallidicinctus",
    "bartramia longicauda",
    "dolichonyx oryzivorus",
    "sturnella magna",
    "sturnella neglecta",
    "circus hudsonius",
    "asio flammeus",
}


class ScoredSpecies(TypedDict):
    scientific_name: str
    common_name: str
    count: int
    score: float
    group: str | None
    context: dict[str, Any]


def get_species_image(scientific_name: str) -> str | None:
    try:
        title = scientific_name.replace(" ", "_")
        url = f"https://en.wikipedia.org/api/rest_v1/page/summary/{title}"

        headers = {
            "User-Agent": "sound-detection/0.2 (https://github.com/dcaulton/sound-detection; bioacoustics research)"
        }

        log.info(f"Fetching image for {scientific_name} → {url}")
        resp = requests.get(url, headers=headers, timeout=8)
        log.info(f"  status={resp.status_code}")

        if resp.status_code == 200:
            data = resp.json()
            thumb = str(data.get("thumbnail", {}).get("source"))
            log.info(f"  thumbnail={thumb}")
            return thumb
        else:
            log.warning(f"  non-200 response fetching image for {scientific_name} (status={resp.status_code})")
            # Optional: log a bit of the body for debugging
            # log.warning(f"  body={resp.text[:200]}")
    except Exception as e:
        log.exception(f"Image fetch failed for {scientific_name}: {e}")
    return None


def get_species_group(scientific_name: str) -> str | None:
    """Return the group (owl / raptor / hummingbird) or None."""
    if not scientific_name:
        return None
    genus = scientific_name.split()[0].lower()
    return SPECIES_GROUPS.get(genus)


def score_species(
    scientific_name: str,
    count: int,
    context: dict[str, Any] | None = None,
    current_month: int | None = None,
) -> float:
    """
    Deterministic interestingness score.
    Higher = more worth highlighting in the summary.
    """
    context = context or {}
    name_lower = scientific_name.lower()
    score = 0.0

    if current_month is None:
        current_month = datetime.now().month

    # 1. Frequency (log scale, capped)
    score += min(math.log1p(count) * 2.5, 12.0)

    # 2. Special groups
    group = get_species_group(scientific_name)
    if group == "owl":
        score += 9.0
    elif group == "raptor":
        score += 8.0
    elif group == "hummingbird":
        score += 7.0

    # 3. Explicit high-interest / indicator list
    if name_lower in HIGH_INTEREST_SPECIES:
        score += 14.0

    # 4. Residency / expected presence
    if context.get("recorded_in_illinois") is False:
        score += 6.0

    residency = (context.get("residency_status") or "").lower()
    if any(k in residency for k in ("vagrant", "rare", "accidental")):
        score += 5.0
    if "migrant" in residency:
        score += 2.0

    # 5. Seasonality - unexpected for current month
    peak_months: list[int] = context.get("migration_peak_months") or context.get("peak_months") or []
    if peak_months and current_month not in peak_months:
        # Present outside its normal peak window
        score += 4.0

    # 6. Indicator flag from Neo4j
    if context.get("is_indicator") or context.get("indicator_for_habitats"):
        score += 7.0

    return round(score, 2)


def bucket_species(
    species_counts: dict[str, int],
    species_contexts: dict[str, dict[str, Any]],
    current_month: int | None = None,
) -> dict[str, list[ScoredSpecies]]:
    scored: list[ScoredSpecies] = []
    for name, count in species_counts.items():
        if count <= 1:
            continue

        context = species_contexts.get(name, {})
        score = score_species(name, count, context, current_month)
        common_name = context.get("common_name") or ""

        scored.append(
            {
                "scientific_name": name,
                "common_name": common_name,
                "count": count,
                "score": score,
                "group": get_species_group(name),
                "context": context,
            }
        )

    by_score = sorted(scored, key=lambda x: x["score"], reverse=True)
    by_count = sorted(scored, key=lambda x: x["count"], reverse=True)

    return {
        "table_rows": by_count,
        "high_interest": [s for s in by_score if s["score"] >= 12][:12],
        "raptors_owls_hummingbirds": [s for s in by_score if s["group"] in ("owl", "raptor", "hummingbird")],
        "rare_vagrant": [
            s
            for s in by_score
            if s["count"] <= 4 and (s["context"].get("recorded_in_illinois") is False or s["score"] >= 10)
        ][:10],
        "dominant": by_count[:8],
    }


class BiomeSummaryService:
    def __init__(
        self,
        session: AsyncSession,
        neo4j_driver: Driver | None = None,
    ) -> None:
        self.session = session
        self.neo4j_driver = neo4j_driver
        self.chunks_per_species: int = 10

        # Create SpeciesKnowledgeService only if we have a driver
        self.species_service = SpeciesKnowledgeService(neo4j_driver) if neo4j_driver else None
        self.retriever = Retriever(neo4j_driver) if neo4j_driver else None

    def _build_human_narrative_chain(self) -> Runnable:
        prompt = ChatPromptTemplate.from_template(
            """You are an experienced field naturalist writing a clear, engaging summary for a 
               homeowner or land steward.

    Write a well-structured markdown report using headings and short paragraphs.

    **Structure to follow:**
    1. Overall soundscape - what was actually common and dominant
    2. Raptors, owls and hummingbirds - comment on diversity when present
    3. High-interest and indicator species
    4. Rare or unexpected species
    5. Brief ecological observations and closing
    6. Unconfirmed / single-analyzer detections (brief, cautious)

    **Important:**
    - Balance common/dominant species with rare and high-interest ones.
    - Use both common and scientific names.
    - Keep the tone knowledgeable but accessible. Aim for 700-1000 words.
    - Treat CONFIRMED detections (BirdNET + Perch agreement) as reliable.
    - Treat SPECULATED detections as possible only; do not lead with them.
    - Prefer confirmed data for activity level, dominant species, and ecological claims.

    **Data**
    Time window: last {window_days} days
    Confirmed species: {total_confirmed_species}
    Speculated (single-analyzer) species: {total_speculated_species}

    Confirmed dominant:
    {dominant_text}

    Confirmed high-interest / indicator:
    {high_interest_text}

    Confirmed rare / unexpected:
    {rare_text}

    Confirmed special groups:
    {group_highlight}

    Speculated (mention cautiously, if at all):
    {speculated_text}

    Human-readable report:"""
        )
        return prompt | llm | StrOutputParser()

    async def create_summary_job(self, site_id: UUID, window_days: int = 30) -> UUID:
        """Creates a pending summary job and returns the ID immediately."""
        summary = BiomeSummary(
            site_id=site_id,
            window_days=window_days,
            status="pending",
        )
        self.session.add(summary)
        await self.session.commit()
        await self.session.refresh(summary)

        log.info(f"Created biome summary job {summary.id} for site {site_id}")
        return summary.id

    async def generate_summary(self, summary_id: UUID) -> None:
        summary = await self.session.get(BiomeSummary, summary_id)
        if not summary:
            log.warning(f"Summary {summary_id} not found")
            return

        start_time = datetime.now(UTC)
        log.info(f"[{summary_id}] Starting biome summary generation (window: {summary.window_days} days)")

        try:
            summary.status = "processing"
            await self.session.commit()

            # === Phase 1: Data Gathering & Filtering ===
            log.info(f"[{summary_id}] Phase 1: Gathering detections and filtering species...")
            phase_start = datetime.now(UTC)
            (
                confirmed_counts,
                speculated_counts,
                confirmed_species,
                speculated_species,
            ) = await self._gather_and_filter_species(summary)

            # Optional: still expose a combined total for old consumers
            species_counts = {
                name: confirmed_counts.get(name, 0) + speculated_counts.get(name, 0)
                for name in set(confirmed_counts) | set(speculated_counts)
            }
            log.info(f"[{summary_id}] Phase 1 completed in {(datetime.now(UTC) - phase_start).seconds}s")

            # === Phase 2: Per-Species Enrichment ===
            phase_start = datetime.now(UTC)
            log.info(
                f"[{summary_id}] Phase 2: Enriching {len(confirmed_species)} notable confirmed species "
                + "with RAG + LLM..."
            )
            enriched_confirmed_species = await self._enrich_species(summary_id, confirmed_species, summary.window_days)
            log.info(
                f"[{summary_id}] Phase 2: Enriching {len(speculated_species)} notable speculated species "
                + "with RAG + LLM..."
            )
            enriched_speculated_species = await self._enrich_species(
                summary_id, speculated_species, summary.window_days
            )

            log.info(f"[{summary_id}] Phase 2 completed in {(datetime.now(UTC) - phase_start).seconds}s")

            # === Phase 3: Narrative Generation ===
            log.info(f"[{summary_id}] Phase 3: Generating final narrative...")
            phase_start = datetime.now(UTC)
            machine_narrative, human_narrative, images, species_table = await self._generate_narratives(
                summary_id=summary.id,
                confirmed_counts=confirmed_counts,
                enriched_confirmed_species=enriched_confirmed_species,
                speculated_counts=speculated_counts,
                enriched_speculated_species=enriched_speculated_species,
                window_days=summary.window_days,
            )
            summary.narrative = machine_narrative
            summary.human_narrative = human_narrative
            summary.notable_species_images = images
            summary.species_table = species_table  # type: ignore[assignment]
            log.info(f"[{summary_id}] Phase 3 completed in {(datetime.now(UTC) - phase_start).seconds}s")

            # Save final results
            # TODO add speculated species and speculated species counts
            summary.summary_json = {
                "window_days": summary.window_days,
                "total_detections": sum(species_counts.values()),
                "total_species": len(confirmed_counts | speculated_counts),
                "species_counts": confirmed_counts,
                "notable_species": enriched_confirmed_species,
                "generated_at": datetime.now(UTC).isoformat(),
            }
            summary.status = "completed"

            duration = (datetime.now(UTC) - start_time).seconds
            log.info(f"[{summary_id}] Summary generation completed successfully in {duration}s")

        except Exception as e:
            summary.status = "failed"
            summary.error_message = str(e)
            log.exception(f"[{summary_id}] Summary generation failed")
        finally:
            await self.session.commit()

    async def list_summaries(self, site_id: UUID, limit: int = 20) -> list[dict]:
        stmt = (
            select(BiomeSummary)
            .where(BiomeSummary.site_id == site_id)  # type: ignore[arg-type]
            .order_by(BiomeSummary.created_at.desc())  # type: ignore[attr-defined]
            .limit(limit)
        )
        result = await self.session.execute(stmt)
        summaries = result.scalars().all()

        return [
            {
                "id": str(s.id),
                "created_at": s.created_at.isoformat(),
                "status": s.status,
                "window_days": s.window_days,
            }
            for s in summaries
        ]

    async def get_summary(self, summary_id: UUID) -> dict | None:
        summary = await self.session.get(BiomeSummary, summary_id)
        if not summary:
            return None

        return {
            "id": str(summary.id),
            "site_id": str(summary.site_id),
            "created_at": summary.created_at.isoformat(),
            "status": summary.status,
            "window_days": summary.window_days,
            "summary_json": summary.summary_json,
            "narrative": summary.narrative,  # existing (Weathervane)
            "human_narrative": summary.human_narrative,  # new pretty version
            "notable_species_images": summary.notable_species_images or [],
            "species_table": summary.species_table or [],
            "grok_narrative": summary.grok_narrative,  # optional high-powered version
            "error_message": summary.error_message,
        }

    async def delete_summary(self, summary_id: UUID) -> bool:
        summary = await self.session.get(BiomeSummary, summary_id)
        if not summary:
            return False

        await self.session.delete(summary)
        await self.session.commit()
        return True

    async def _gather_and_filter_species(
        self, summary: BiomeSummary
    ) -> tuple[dict[str, int], dict[str, int], list[dict], list[dict]]:
        """
        Returns:
            confirmed_counts:   scientific_name -> distinct confirmed_group_id count
            speculated_counts:  scientific_name -> single-model detection count
            confirmed_species:  filtered list for enrichment (confirmation=confirmed)
            speculated_species: filtered list for enrichment (confirmation=speculated)
        """
        log.info(f"[{summary.id}] Querying detections from last {summary.window_days} days...")
        cutoff = datetime.now(UTC) - timedelta(days=summary.window_days)

        # --- Confirmed: agreement events (distinct group ids per species) ---
        confirmed_stmt = (
            select(  # type: ignore[call-overload]
                Detection.scientific_name,
                func.count(func.distinct(Detection.confirmed_group_id)).label("count"),  # type: ignore[arg-type]
            )  # type: ignore[arg-type]
            .join(Recording, Detection.recording_id == Recording.id)
            .join(Microphone, Recording.microphone_id == Microphone.id)
            .where(Microphone.site_id == summary.site_id)
            .where(Detection.created_at >= cutoff)
            .where(col(Detection.confirmed_group_id).is_not(None))
            .group_by(Detection.scientific_name)
            .order_by(func.count(func.distinct(Detection.confirmed_group_id)).desc())  # type: ignore[arg-type]
        )

        # --- Speculated: single-model rows only ---
        speculated_stmt = (
            select(  # type: ignore[call-overload]
                Detection.scientific_name,
                func.count().label("count"),  # type: ignore[arg-type]
            )  # type: ignore[arg-type]
            .join(Recording, Detection.recording_id == Recording.id)
            .join(Microphone, Recording.microphone_id == Microphone.id)
            .where(Microphone.site_id == summary.site_id)
            .where(Detection.created_at >= cutoff)
            .where(col(Detection.confirmed_group_id).is_(None))
            .group_by(Detection.scientific_name)
            .order_by(func.count().desc())  # type: ignore[arg-type]
        )

        confirmed_rows = (await self.session.execute(confirmed_stmt)).all()
        speculated_rows = (await self.session.execute(speculated_stmt)).all()

        confirmed_counts: dict[str, int] = {}
        for row in confirmed_rows:
            if row.scientific_name:
                confirmed_counts[row.scientific_name] = int(row.count)  # type: ignore[call-overload]

        speculated_counts: dict[str, int] = {}
        for row in speculated_rows:
            if row.scientific_name:
                speculated_counts[row.scientific_name] = int(row.count)  # type: ignore[call-overload]

        if not confirmed_counts and not speculated_counts:
            log.info(f"[{summary.id}] No detections found in window.")
            return {}, {}, [], []

        all_counts = {**speculated_counts, **confirmed_counts}  # confirmed overwrites if both
        max_count = max(all_counts.values())
        log.info(
            f"[{summary.id}] Found {len(confirmed_counts)} confirmed, "
            f"{len(speculated_counts)} speculated species. Max count = {max_count}"
        )

        # Get species context from Neo4j (union of names)
        species_contexts: dict[str, dict] = {}
        if self.species_service is not None:
            for name in set(confirmed_counts) | set(speculated_counts):
                context = self.species_service.get_species_by_scientific_name(name)
                if context:
                    species_contexts[name] = context

        def _passes_filters(name: str, count: int, *, drop_singles: bool) -> bool:
            name_lower = name.lower().strip()

            # 1. Skip known non-species labels
            if name_lower in NON_SPECIES_LABELS:
                return False

            # 2. Skip single-observation detections (likely false positives)
            #    Confirmed events skip this: one agreement is already strong signal
            if drop_singles and count <= 1:
                return False

            # 3. Require it to look roughly like a scientific name
            if " " not in name or not name[0].isupper():
                return False

            # this is a little too agressive for now, and messes up multitenancy.
            # TODO: have a user or site setting drive this later
            # context = species_contexts.get(name, {})
            # recorded_in_illinois = context.get("recorded_in_illinois", False)
            # if not recorded_in_illinois:
            #     return False

            return True

        confirmed_species: list[dict] = []
        for name, count in confirmed_counts.items():
            if not _passes_filters(name, count, drop_singles=False):
                continue
            confirmed_species.append(
                {
                    "scientific_name": name,
                    "count": count,
                    "confirmation": "confirmed",
                    "context": species_contexts.get(name, {}),
                }
            )

        speculated_species: list[dict] = []
        for name, count in speculated_counts.items():
            if not _passes_filters(name, count, drop_singles=True):
                continue
            speculated_species.append(
                {
                    "scientific_name": name,
                    "count": count,
                    "confirmation": "speculated",
                    "context": species_contexts.get(name, {}),
                }
            )

        log.info(
            f"[{summary.id}] After filtering: "
            f"confirmed={len(confirmed_species)} "
            f"(from {len(confirmed_counts)}), "
            f"speculated={len(speculated_species)} "
            f"(from {len(speculated_counts)})"
        )

        return confirmed_counts, speculated_counts, confirmed_species, speculated_species

    def _build_species_enrichment_chain(self) -> Runnable:
        prompt = ChatPromptTemplate.from_template(
            """You are an expert ornithologist and ecologist analyzing bird observations.

    Species: {scientific_name} ({common_name})
    Detection count in the last {window_days} days: {count}

    Structured knowledge:
    - Recorded in Illinois: {recorded_in_illinois}
    - Primary habitat: {primary_habitat}
    - Diet: {diet_description}
    - Interesting facts: {interesting_facts}

    Relevant context from Wikipedia:
    {chunks_text}

    Task:
    Write a concise but insightful paragraph (3-6 sentences) that explains:
    - Why this species might be notable or unusual in this location right now
    - Any relevant seasonal, migratory, or ecological context
    - Interesting biological or behavioral details worth mentioning

    Be factual and avoid speculation. If the species is common and expected, say so briefly.
    """
        )

        chain = prompt | llm | StrOutputParser()
        return chain

    async def _enrich_species(self, summary_id: UUID, filtered_species: list[dict], window_days: int) -> list[dict]:
        if not self.retriever:
            log.warning(f"[{summary_id}] No retriever available - skipping RAG step")
            return filtered_species

        log.info(f"[{summary_id}] Enriching {len(filtered_species)} species using LangChain + RAG...")

        chain = self._build_species_enrichment_chain()
        enriched = []

        for species in filtered_species:
            name = species["scientific_name"]
            count = species["count"]
            context = species.get("context", {})

            log.info(f"[{summary_id}] → Enriching {name} ({count} detections)")

            # Get top relevant chunks
            chunks = self.retriever.retrieve(name, top_k=self.chunks_per_species)

            # Format chunks for the prompt
            chunks_text = (
                "\n\n".join(f"- {chunk.get('text', '')[:800]}" for chunk in chunks[:5])
                if chunks
                else "No additional context available."
            )

            try:
                insight = await chain.ainvoke(
                    {
                        "scientific_name": name,
                        "common_name": context.get("common_name", name),
                        "window_days": window_days,
                        "count": count,
                        "recorded_in_illinois": context.get("recorded_in_illinois", False),
                        "primary_habitat": context.get("primary_habitat", "Unknown"),
                        "diet_description": context.get("diet_description", "Unknown"),
                        "interesting_facts": context.get("interesting_facts", "None available"),
                        "chunks_text": chunks_text,
                    }
                )

                enriched.append(
                    {
                        **species,
                        "chunks": chunks,
                        "llm_insight": insight.strip(),
                    }
                )

            except Exception as e:
                log.warning(f"[{summary_id}] LLM enrichment failed for {name}: {e}")
                enriched.append(
                    {
                        **species,
                        "chunks": chunks,
                        "llm_insight": None,
                    }
                )

        log.info(f"[{summary_id}] Finished enriching {len(enriched)} species with LangChain")
        return enriched

    def _build_narrative_chain(self) -> Runnable:
        prompt = ChatPromptTemplate.from_template(
            """You are an expert ornithologist and ecologist producing a concise, structured summary of 
           acoustic detections for downstream automated systems (Weathervane and similar tools).

Write a clear, information-dense report. Use short paragraphs and bullet points where helpful. Avoid fluff.

**Required content:**
- Overall activity level and the dominant (most frequently detected) species
- Any notable diversity in owls, raptors, or hummingbirds
- High-interest or indicator species that stand out
- Brief mention of rare or unexpected species
- One or two ecological observations worth tracking

**Important:**
- Treat CONFIRMED detections (BirdNET + Perch agreement) as reliable.
- Treat SPECULATED detections as possible only; do not lead with them.
- Prefer confirmed data for activity level, dominant species, and ecological claims.

Time window: last {window_days} days
Confirmed species: {total_confirmed_species}
Speculated (single-analyzer) species: {total_speculated_species}

Confirmed dominant:
{dominant_text}

Confirmed high-interest / indicator:
{high_interest_text}

Confirmed rare / unexpected:
{rare_text}

Confirmed special groups:
{group_highlight}

Speculated (mention cautiously):
{speculated_text}

Structured summary:"""
        )
        return prompt | llm | StrOutputParser()

    def _add_species_table(self, doc: DocumentObject, table_rows: list[ScoredSpecies]) -> None:
        """Add a properly formatted Word table."""
        if not table_rows:
            doc.add_paragraph("No species with more than one detection.")
            return

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        # Header
        headers = ["Common Name", "Scientific Name", "Detections", "Score"]
        header_cells = table.rows[0].cells
        for i, text in enumerate(headers):
            header_cells[i].text = text
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Data rows
        for row in table_rows:
            cells = table.add_row().cells
            cells[0].text = row.get("common_name") or "—"
            cells[1].text = row.get("scientific_name") or ""
            cells[2].text = str(row.get("count", ""))
            cells[3].text = str(row.get("score", ""))

        doc.add_paragraph()  # spacer

    async def export_to_docx(self, summary: dict) -> Path:
        """
        Generate a nicely formatted Word document from a completed summary.
        Returns the path to a temporary .docx file.
        """
        doc = Document()

        # ---------- Title ----------
        title = doc.add_heading("Biome Summary Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ---------- Metadata ----------
        meta = doc.add_paragraph()
        meta.add_run("Site ID: ").bold = True
        meta.add_run(str(summary.get("site_id", "Unknown")))
        meta.add_run("\n")
        meta.add_run("Window: ").bold = True
        meta.add_run(f"Last {summary.get('window_days', '?')} days")
        meta.add_run("\n")
        meta.add_run("Generated: ").bold = True
        created = summary.get("created_at", "")
        meta.add_run(created[:19].replace("T", " ") if created else "Unknown")

        doc.add_paragraph()  # spacer

        # ---------- Human Narrative ----------
        doc.add_heading("Summary", level=1)

        human_narrative = summary.get("human_narrative") or summary.get("narrative") or "No narrative available."

        # Simple paragraph splitting (works well enough for markdown-ish text)
        for block in human_narrative.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if block.startswith("#"):
                # Treat as a heading
                level = min(block.count("#"), 3)
                text = block.lstrip("# ").strip()
                doc.add_heading(text, level=level)
            else:
                doc.add_paragraph(block)

        # ----- Species Table -----
        doc.add_heading("Species Detected (more than one detection)", level=1)

        table_rows = summary.get("species_table") or []
        self._add_species_table(doc, table_rows)

        # ----- Notable Species with Images -----
        images = summary.get("notable_species_images") or []
        if images:
            doc.add_heading("Notable Species", level=1)

            headers = {
                "User-Agent": "sound-detection/0.2 (https://github.com/dcaulton/sound-detection; bioacoustics research)"
            }

            for item in images:
                name = item.get("scientific_name", "Unknown")
                common = item.get("common_name") or ""
                count = item.get("count", "?")
                image_url = item.get("image_url")

                caption = f"{common} ({name})" if common else name
                caption += f"  —  {count} detections"
                doc.add_heading(caption, level=2)

                if image_url:
                    try:
                        resp = requests.get(image_url, headers=headers, timeout=10)
                        if resp.status_code != 200:
                            doc.add_paragraph(f"[Image unavailable - HTTP {resp.status_code}]")
                            continue

                        # Convert to a format python-docx understands (JPEG)
                        image_stream = BytesIO(resp.content)
                        with PILImage.open(image_stream) as img:
                            # Convert palette/RGBA images etc. to RGB
                            if img.mode in ("RGBA", "P", "LA"):
                                img = img.convert("RGB")  # type: ignore[assignment]
                            elif img.mode != "RGB":
                                img = img.convert("RGB")  # type: ignore[assignment]

                            output = BytesIO()
                            img.save(output, format="JPEG", quality=85)
                            output.seek(0)

                            doc.add_picture(output, width=Inches(3.5))
                            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    except Exception as e:
                        doc.add_paragraph(f"[Could not load image: {type(e).__name__}: {e}]")
                        log.exception(f"Failed to embed image {image_url}")

                doc.add_paragraph()

        # ---------- Footer ----------
        doc.add_paragraph()
        footer = doc.add_paragraph("Generated by sound-detection")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(120, 120, 120)

        # Save to a temporary file
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()

        return Path(tmp.name)

    async def _generate_narratives(
        self,
        summary_id: UUID,
        confirmed_counts: dict[str, int],
        enriched_confirmed_species: list[dict],
        speculated_counts: dict[str, int],
        enriched_speculated_species: list[dict],
        window_days: int,
    ) -> tuple[str, str, list[dict[str, Any]], list[ScoredSpecies]]:
        """
        Returns:
            (machine_narrative, human_narrative, notable_species_images, species_table)
        """
        log.info(f"[{summary_id}] Generating narratives...")

        # ----- Context from enrichment (confirmed preferred, then speculated) -----
        species_contexts: dict[str, dict] = {}
        for sp in enriched_confirmed_species + enriched_speculated_species:
            name = sp["scientific_name"]
            species_contexts[name] = sp.get("context") or {}

        for name in set(confirmed_counts) | set(speculated_counts):
            species_contexts.setdefault(name, {})

        # ----- Score & bucket on CONFIRMED only (hard-won signal) -----
        buckets = bucket_species(confirmed_counts, species_contexts)

        # Speculated: score for a short secondary list (still use same scorer)
        speculated_buckets = bucket_species(speculated_counts, species_contexts)

        # ----- Group highlight (confirmed) -----
        group_parts = []
        for group_key, label in [
            ("owl", "Owls"),
            ("raptor", "Raptors"),
            ("hummingbird", "Hummingbirds"),
        ]:
            members = [r for r in buckets["raptors_owls_hummingbirds"] if r["group"] == group_key]
            if not members:
                continue
            names = []
            for m in members:
                common = m["common_name"] or ""
                sci = m["scientific_name"]
                if common:
                    names.append(f"{common} (*{sci}*)")
                else:
                    names.append(f"*{sci}*")
            group_parts.append(f"{label}: {len(members)} species ({', '.join(names)})")

        group_highlight = (
            "\n".join(group_parts)
            if group_parts
            else "No major raptor, owl, or hummingbird diversity noted among confirmed detections."
        )

        def format_species_list(rows: list[ScoredSpecies], limit: int = 8) -> str:
            parts = []
            for r in rows[:limit]:
                common = r["common_name"] or ""
                label = f"{common} (*{r['scientific_name']}*)" if common else f"*{r['scientific_name']}*"
                parts.append(f"- {label}: {r['count']} events (score {r['score']})")
            return "\n".join(parts) if parts else "None"

        high_interest_text = format_species_list(buckets["high_interest"])
        dominant_text = format_species_list(buckets["dominant"])
        rare_text = format_species_list(buckets["rare_vagrant"])
        # Speculated: top by score only, capped, clearly labeled later in prompts
        speculated_text = format_species_list(
            speculated_buckets["high_interest"] or speculated_buckets["dominant"],
            limit=6,
        )

        prompt_vars = {
            "window_days": window_days,
            "total_confirmed_species": len(confirmed_counts),
            "total_speculated_species": len(speculated_counts),
            "group_highlight": group_highlight,
            "high_interest_text": high_interest_text,
            "dominant_text": dominant_text,
            "rare_text": rare_text,
            "speculated_text": speculated_text,
        }

        machine_chain = self._build_narrative_chain()
        machine_narrative = await machine_chain.ainvoke(prompt_vars)

        human_chain = self._build_human_narrative_chain()
        human_narrative = await human_chain.ainvoke(prompt_vars)

        # Images: prefer confirmed high-interest
        notable_species_images: list[dict[str, Any]] = []
        image_candidates: list[Any] = buckets["high_interest"] or buckets["dominant"]
        for sp in image_candidates[:5]:
            name = sp["scientific_name"]
            image_url = get_species_image(name)
            if image_url:
                notable_species_images.append(
                    {
                        "scientific_name": name,
                        "common_name": sp.get("common_name"),
                        "count": sp["count"],
                        "confirmation": "confirmed",
                        "image_url": image_url,
                    }
                )

        # Table: confirmed rows first (by count), then mark confirmation in consumer if needed
        return (
            machine_narrative.strip(),
            human_narrative.strip(),
            notable_species_images,
            buckets["table_rows"],
        )

    async def build_grok_data_package(self, summary_id: UUID) -> str | None:
        """
        Build a clean, paste-ready markdown package of the summary data
        for use with a stronger external model (Grok, etc.).
        """
        summary = await self.session.get(BiomeSummary, summary_id)
        if not summary or summary.status != "completed":
            return None

        table: list[dict] = summary.species_table or []
        if not table:
            return "No species table available for this summary."

        # Re-derive the useful buckets from the stored table
        by_score = sorted(table, key=lambda x: x.get("score", 0), reverse=True)
        by_count = sorted(table, key=lambda x: x.get("count", 0), reverse=True)

        dominant = by_count[:8]
        high_interest = [s for s in by_score if s.get("score", 0) >= 12][:10]
        group_members = [s for s in by_score if s.get("group") in ("owl", "raptor", "hummingbird")]
        rare = [s for s in by_score if s.get("count", 0) <= 4 and s.get("score", 0) >= 10][:8]

        def fmt(rows: list[dict]) -> str:
            if not rows:
                return "_None_"
            lines = []
            for r in rows:
                common = r.get("common_name") or "—"
                sci = r.get("scientific_name", "")
                count = r.get("count", "?")
                score = r.get("score", "?")
                lines.append(f"- {common} (*{sci}*): {count} detections (score {score})")
            return "\n".join(lines)

        # Group highlight
        group_lines = []
        for key, label in [("owl", "Owls"), ("raptor", "Raptors"), ("hummingbird", "Hummingbirds")]:
            members = [s for s in group_members if s.get("group") == key]
            if members:
                names = [
                    f"{s.get('common_name') or s.get('scientific_name')} (*{s.get('scientific_name')}*)"
                    for s in members
                ]
                group_lines.append(f"- {label} ({len(members)}): {', '.join(names)}")
        group_highlight = "\n".join(group_lines) if group_lines else "- None noted"

        package = f"""# Biome Summary Data Package
    Site ID: {summary.site_id}
    Window: last {summary.window_days} days
    Generated: {summary.created_at.isoformat() if summary.created_at else "unknown"}
    Total species with >1 detection: {len(table)}

    ## Dominant / Most Frequent Species
    {fmt(dominant)}

    ## High-Interest & Indicator Species
    {fmt(high_interest)}

    ## Raptors, Owls & Hummingbirds
    {group_highlight}

    ## Rare / Unexpected Species
    {fmt(rare)}

    ## Full Species Table (count > 1)
    | Common Name | Scientific Name | Detections | Score |
    |-------------|-----------------|------------|-------|
    """
        for r in by_count:
            common = r.get("common_name") or "—"
            sci = r.get("scientific_name", "")
            package += f"| {common} | {sci} | {r.get('count', '')} | {r.get('score', '')} |\n"

        package += """
    ---
    Instructions for the model:
    You are an experienced field naturalist and ecologist. Using only the data above, write a clear, engaging, 
    well-structured markdown report for a homeowner / land steward in the Chicago suburbs / prairie restoration 
    context. Balance the common soundscape with high-interest, indicator, raptor/owl, and rare species. Use 
    both common and scientific names. Aim for 700-1000 words.
    """
        return package

    async def update_grok_narrative(self, summary_id: UUID, text: str) -> bool:
        summary = await self.session.get(BiomeSummary, summary_id)
        if not summary:
            return False

        summary.grok_narrative = text
        self.session.add(summary)
        await self.session.commit()
        return True

    async def backfill_images(self, summary_id: UUID) -> dict | None:
        summary = await self.session.get(BiomeSummary, summary_id)
        if not summary:
            return None

        candidates = summary.species_table or []
        candidates = sorted(
            candidates,
            key=lambda x: x.get("score", 0),
            reverse=True,
        )[:8]

        images = []
        for sp in candidates:
            name = sp.get("scientific_name")
            if not name:
                continue
            image_url = get_species_image(name)
            if image_url:
                images.append(
                    {
                        "scientific_name": name,
                        "common_name": sp.get("common_name"),
                        "count": sp.get("count"),
                        "image_url": image_url,
                    }
                )

        summary.notable_species_images = images
        self.session.add(summary)
        await self.session.commit()
        await self.session.refresh(summary)

        return {
            "summary_id": str(summary.id),
            "images_found": len(images),
            "notable_species_images": images,
        }
